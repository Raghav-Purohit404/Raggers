import sys
import os
import pandas as pd
import streamlit as st
import tempfile
from pathlib import Path
from datetime import datetime
import time

# ✅ Set Streamlit page config FIRST
st.set_page_config(page_title="PhiRAG: Chat with Your Knowledge", layout="wide")

# ─────────────────────────────────────────────────────────────
#  PATH & IMPORT SETUP
# ─────────────────────────────────────────────────────────────

# Add root repo to sys.path for module imports
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from ingestion import (
    load_documents_from_files,
    load_documents_from_urls,
    get_vectorstore,
    sync_to_backend_faiss  # 🔁 Incremental FAISS sync
)
from logger import log_query
from llm_wrapper import get_llm_response  # ⬅️ use get_llm_response from wrapper
from rag_pipeline import run_pipeline  # fallback LLM pipeline

# Define paths
BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
LOG_DIR = os.path.join(BASE_DIR, "logs")
INDEX_PATH = os.path.join(BASE_DIR, "combined_faiss_index")
LOG_PATH = os.path.join(LOG_DIR, "query_logs.csv")

# ─────────────────────────────────────────────────────────────
# ✅ Attempt to auto-load FAISS index on page load
# ─────────────────────────────────────────────────────────────
if "vectorstore_ready" not in st.session_state:
    if os.path.exists(INDEX_PATH):
        try:
            db = get_vectorstore([], rebuild=False, load_path=INDEX_PATH)
            st.session_state["vectorstore_ready"] = True
            st.success("✅ FAISS index auto-loaded.")
        except Exception as e:
            st.warning(f"⚠️ Error loading FAISS index: {e}")
            st.session_state["vectorstore_ready"] = False
    else:
        st.session_state["vectorstore_ready"] = False

# Ensure logging directory and file exist
os.makedirs(LOG_DIR, exist_ok=True)
if not os.path.exists(LOG_PATH) or os.path.getsize(LOG_PATH) == 0:
    pd.DataFrame(columns=["Timestamp", "Query", "Response", "Feedback"]).to_csv(LOG_PATH, index=False)

# ─────────────────────────────────────────────────────────────
# PAGE LAYOUT
# ─────────────────────────────────────────────────────────────
st.title("📚 PhiRAG: Chat with Files + Web + LLM")

# ─────────────────────────────────────────────────────────────
# INGESTION SECTION
# ─────────────────────────────────────────────────────────────

uploaded_files = st.file_uploader(
    "Upload PDF, TXT, DOCX, CSV, or MD files",
    type=["pdf", "txt", "docx", "csv", "md"],
    accept_multiple_files=True,
)

folder_path = st.text_input("📁 Or enter a local folder path:")
url_input = st.text_area("🌐 Paste Web URLs (one per line):")
rebuild = st.checkbox("🔄 Force rebuild FAISS index")

if "frontend_docs" not in st.session_state:
    st.session_state["frontend_docs"] = []

process = st.button("📥 Ingest Files and Links")

@st.cache_data(show_spinner=False)
def save_uploaded_files(uploaded_files):
    temp_dir = tempfile.mkdtemp()
    paths = []
    for file in uploaded_files:
        path = os.path.join(temp_dir, file.name)
        with open(path, "wb") as f:
            f.write(file.read())
        paths.append(path)
    return paths

@st.cache_data(show_spinner=False)
def summarize_file(path):
    try:
        ext = Path(path).suffix.lower()
        if ext in [".txt", ".md"]:
            with open(path, "r", encoding="utf-8") as f:
                lines = f.readlines()
            return f"Lines: {len(lines)}, Preview: {lines[0][:100] if lines else 'Empty'}"
        elif ext == ".csv":
            df = pd.read_csv(path)
            return f"Rows: {len(df)}, Columns: {len(df.columns)}"
        elif ext == ".docx":
            import docx
            doc = docx.Document(path)
            return f"Paragraphs: {len(doc.paragraphs)}, First: {doc.paragraphs[0].text[:100] if doc.paragraphs else 'Empty'}"
        elif ext == ".pdf":
            import fitz
            doc = fitz.open(path)
            return f"Pages: {len(doc)}, Preview: {doc[0].get_text()[:100] if len(doc) > 0 else 'Empty'}"
    except Exception as e:
        return f"Error: {e}"

if process:
    file_paths = []

    if uploaded_files:
        file_paths.extend(save_uploaded_files(uploaded_files))

    if folder_path and os.path.exists(folder_path):
        for fname in os.listdir(folder_path):
            if fname.lower().endswith((".pdf", ".txt", ".docx", ".csv", ".md")):
                file_paths.append(os.path.join(folder_path, fname))

    urls = url_input.strip().splitlines() if url_input.strip() else []

    st.subheader("📋 File Summary Preview")
    for f in file_paths:
        st.markdown(f"**{os.path.basename(f)}**")
        st.text(summarize_file(f))

    documents = load_documents_from_files(file_paths) + load_documents_from_urls(urls)
    st.session_state["frontend_docs"] = documents  # 💾 Save for syncing later

    if rebuild:
        if not documents:
            st.error("❌ Upload files or provide valid URLs before rebuilding FAISS index.")
        else:
            db = get_vectorstore(documents, rebuild=True, save_path=INDEX_PATH)
            st.success("✅ FAISS index rebuilt.")
            st.session_state["vectorstore_ready"] = True
    else:
        if os.path.exists(INDEX_PATH):
            db = get_vectorstore([], rebuild=False, load_path=INDEX_PATH)
            st.success("✅ Loaded existing FAISS index.")
            st.session_state["vectorstore_ready"] = True
        else:
            st.warning("⚠️ No saved FAISS index found. Please ingest documents.")
            st.session_state["vectorstore_ready"] = False

# ─────────────────────────────────────────────────────────────
#  QUERY SECTION
# ─────────────────────────────────────────────────────────────

query = st.text_input("💬 Ask a question:")

# OPTIONAL: Query size selection
st.markdown("### 🧠 (Optional) Choose answer depth:")
col1, col2, col3, col4 = st.columns(4)
answer_type = None

if col1.button("Summary (100 words)"):
    answer_type = "summary"
elif col2.button("Overview (200 words)"):
    answer_type = "overview"
elif col3.button("Detailed (400 words)"):
    answer_type = "detailed"
elif col4.button("Deep Dive (600+ words)"):
    answer_type = "deep_dive"

def get_word_limit(answer_type):
    return {
        "summary": 100,
        "overview": 200,
        "detailed": 400,
        "deep_dive": 600
    }.get(answer_type, 150)

run_query = st.button("🔍 Run Query")

if run_query and query:
    if os.path.exists(INDEX_PATH) and st.session_state.get("vectorstore_ready", False):
        db = get_vectorstore([], rebuild=False, load_path=INDEX_PATH)
        retriever = db.as_retriever(search_kwargs={"k": 5})
        docs = retriever.get_relevant_documents(query)
        context = "\n\n".join(doc.page_content for doc in docs[:5])

        word_limit = get_word_limit(answer_type)
        prompt = f"Context:\n{context}\n\nQuestion: {query}\n\nStrictly answer in exactly {word_limit} words. Count your words."

        # ⏱️ Start timing
        start_time = time.time()
        answer = get_llm_response(prompt, word_limit)
        end_time = time.time()
        response_time = round(end_time - start_time, 2)

        total_file_size_bytes = sum(len(doc.page_content.encode('utf-8')) for doc in docs)
        total_file_size_mb = round(total_file_size_bytes / (1024 * 1024), 2)
        response_size_mb = round(len(answer.encode('utf-8')) / (1024 * 1024), 4)

        st.subheader("💬 Answer")
        st.write(answer)

        with st.expander("📊 Response Metrics"):
            st.markdown(f"**LLM Response Time:** `{response_time}` seconds")
            st.markdown(f"**Total Size of Retrieved Documents:** `{total_file_size_mb}` MB")
            st.markdown(f"**Size of Generated Response:** `{response_size_mb}` MB")

        st.subheader("📌 Source Snippets")
        for i, doc in enumerate(docs, start=1):
            st.markdown(f"**Document {i}: {os.path.basename(doc.metadata.get('source', 'Unknown'))}**")
            st.write(doc.page_content[:500] + "...")

        log_query(query, answer)

        if st.session_state.get("frontend_docs"):
            sync_to_backend_faiss(st.session_state["frontend_docs"], backend_path="faiss_backend")
            st.session_state["frontend_docs"] = []

    else:
        st.warning("⚠️ No FAISS index found. Using LLM-only mode.")
        try:
            start_time = time.time()
            result = run_pipeline(prompt=query)
            end_time = time.time()
            response_time = round(end_time - start_time, 2)
            response_size_mb = round(len(result.encode("utf-8")) / (1024 * 1024), 4)

            st.subheader("💬 Answer (LLM Only)")
            st.write(result)

            with st.expander("📊 Response Metrics"):
                st.markdown(f"**LLM Response Time:** `{response_time}` seconds")
                st.markdown(f"**Size of Generated Response:** `{response_size_mb}` MB")

            log_query(query, result)

        except Exception as e:
            st.error(f"Error: {e}")

# ─────────────────────────────────────────────────────────────
# LOGGING UI
# ─────────────────────────────────────────────────────────────

df_logs = pd.read_csv(LOG_PATH)
if "Feedback" not in df_logs.columns:
    df_logs["Feedback"] = ""

if not df_logs.empty:
    df_logs["Timestamp"] = pd.to_datetime(df_logs["Timestamp"], utc=True, errors='coerce')
    df_logs.sort_values("Timestamp", inplace=True)

st.markdown("---")
st.subheader("🗂️ Query Log Viewer & Export")

search_term = st.text_input("🔎 Filter logs by keyword (in query or response):")

mode = st.radio("📤 Select log filter:", [
    "Download all logs",
    "Download latest log",
    "Download last N logs",
    "Download logs by date range"
])

data_to_download = df_logs.copy()

if mode == "Download latest log":
    data_to_download = df_logs.tail(1)
elif mode == "Download last N logs":
    n = st.number_input("Enter number of recent logs:", min_value=1, max_value=len(df_logs), step=1)
    data_to_download = df_logs.tail(n)
elif mode == "Download logs by date range":
    min_date = df_logs["Timestamp"].min().date()
    max_date = df_logs["Timestamp"].max().date()
    start_date = st.date_input("Start date", value=min_date)
    end_date = st.date_input("End date", value=max_date)
    mask = (df_logs["Timestamp"].dt.date >= start_date) & (df_logs["Timestamp"].dt.date <= end_date)
    data_to_download = df_logs.loc[mask]

if search_term.strip() and not data_to_download.empty:
    keyword = search_term.lower()
    data_to_download = data_to_download[
        data_to_download["Query"].str.lower().str.contains(keyword, na=False) |
        data_to_download["Response"].str.lower().str.contains(keyword, na=False)
    ]

if data_to_download.empty:
    st.warning("No logs match the selected filters.")
else:
    st.markdown("### 🖋️ Feedback and Preview")
    feedback_updates = []

    for i, row in data_to_download.iterrows():
        st.markdown(f"**Timestamp:** {row['Timestamp']}")
        st.markdown(f"**Query:** {row['Query']}")
        st.markdown(f"**Response:** {row['Response']}")
        feedback = st.radio(
            f"Feedback for {row['Timestamp']}",
            options=["", "👍", "👎"],
            horizontal=True,
            key=f"feedback_{i}",
            index=["", "👍", "👎"].index(str(row.get("Feedback", ""))) if row.get("Feedback", "") in ["👍", "👎"] else 0
        )
        feedback_updates.append((row.name, feedback))
        st.markdown("---")

    if st.button("💾 Save Feedback"):
        for idx, fb in feedback_updates:
            df_logs.at[idx, "Feedback"] = fb
        df_logs.to_csv(LOG_PATH, index=False)
        st.success("✅ Feedback saved.")

    file_suffix = mode.lower().replace(" ", "_")
    file_name = f"query_logs_{file_suffix}.csv"
    csv_data = data_to_download.to_csv(index=False).encode("utf-8")
    st.download_button(
        label="📥 Download Filtered Logs as CSV",
        data=csv_data,
        file_name=file_name,
        mime="text/csv"
    )
